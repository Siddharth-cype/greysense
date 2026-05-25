from docx import Document

d = Document(r'C:\Users\Asus\Downloads\hexhiveint-20260206T082306Z-1-001\Siddharth_K_Resume.docx')
for i, p in enumerate(d.paragraphs):
    runs_info = []
    for r in p.runs:
        font_size = r.font.size.pt if r.font.size else 'inherit'
        bold = r.font.bold
        runs_info.append(f'[{font_size}pt, bold={bold}]')
    j = " ".join(runs_info)
    print(f'P{i}: style={p.style.name} | align={p.alignment} | runs={len(p.runs)} | {j}')
print('---TABLES---')
for ti, table in enumerate(d.tables):
    print(f'Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows):
        cells = [cell.text[:50] for cell in row.cells]
        print(f'  Row {ri}: {cells}')
