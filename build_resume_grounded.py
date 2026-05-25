from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== PAGE MARGINS 0.5" ==========
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# ========== HELPER FUNCTIONS ==========
def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_entry_line(left_bold, left_normal, right_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    if left_bold:
        run_b = p.add_run(left_bold)
        run_b.font.size = Pt(12)
        run_b.font.bold = True
    if left_normal:
        run_n = p.add_run(left_normal)
        run_n.font.size = Pt(12)
    if right_text:
        run_t = p.add_run('\t' + right_text)
        run_t.font.size = Pt(12)

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"- {text}")
    run.font.size = Pt(12)

# ========== NAME & CONTACT ==========
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(0)
name_run = name_p.add_run('SIDDHARTH K')
name_run.font.size = Pt(20)
name_run.font.bold = True

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(4)
contact_run = contact_p.add_run('+91 9994475724  |  siddharthkadiresan21@gmail.com  |  LinkedIn: Siddharth K  |  GitHub: Siddharth-cype')
contact_run.font.size = Pt(12)

# ========== SUMMARY ==========
add_section_header('Summary')
summary_p = doc.add_paragraph()
summary_run = summary_p.add_run(
    'Pre-final year B.Tech ECE student with hands-on experience in software development, '
    'embedded systems, IoT, and machine learning. Proficient in Java, Python, C/C++, '
    'and cloud technologies. Experienced in building cloud-connected systems using AWS and ESP32.'
)
summary_run.font.size = Pt(12)

# ========== EDUCATION ==========
add_section_header('Education')
add_entry_line('Manakula Vinayagar Institute of Technology', ' — B.Tech ECE (CGPA: 8.1)', 'Sep 2023 – Apr 2027 | Puducherry, India')
add_entry_line('Vivekananda Higher Secondary School', ' — Higher Secondary Class XII (77.3%)', 'Mar 2023 | Puducherry, India')
add_entry_line('Blessed Mother Teresa Model Higher Secondary School', ' — Secondary Class X (SSLC)', 'Mar 2021 | Puducherry, India')

# ========== TECHNICAL SKILLS ==========
add_section_header('Technical Skills')
skills = [
    ('Languages: ', 'Java, Python, C, C++, Embedded C'),
    ('Embedded & IoT: ', 'ESP32'),
    ('Cloud: ', 'AWS EC2, AWS S3'),
    ('ML & AI: ', 'TensorFlow, Keras, YOLO, Vision Transformer (ViT), OpenCV'),
    ('Databases: ', 'MySQL')
]
for label, value in skills:
    p = doc.add_paragraph()
    run_label = p.add_run(label)
    run_label.font.size = Pt(12)
    run_label.font.bold = True
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)

# ========== PROJECTS ==========
add_section_header('Projects')

add_entry_line('AWS-Based IoT Environmental & Activity Monitoring System', ' (ESP32, AWS, MQTT, Python)', 'Jan 2026')
add_bullet('Programmed an ESP32 multi-sensor node using C/C++ to read environment data, enabling real-time detection of human presence, noise levels, and air quality.')
add_bullet('Built a data pipeline publishing MQTT telemetry to AWS services, allowing web dashboard visualization with low-latency updates.')

add_entry_line('Vision Transformer (ViT) for Image Classification', ' (Python, TensorFlow, Keras)', 'Dec 2025')
add_bullet('Trained a Vision Transformer model on the CIFAR-10 dataset by implementing patch embedding and multi-head attention layers from scratch.')
add_bullet('Evaluated classification performance using accuracy metrics against test datasets, achieving reliable detection of unseen objects.')

add_entry_line('Computer Vision-Based Touchless System Control', ' (Python, OpenCV)', 'May 2025')
add_bullet('Developed a hand gesture recognition script using OpenCV to track finger movements via the laptop camera.')
add_bullet('Mapped detected hand coordinates to operating system controls, enabling full touchless mouse operation including cursor movement and clicking.')

# ========== INTERNSHIP ==========
add_section_header('Internship')
add_entry_line('Software Engineering Intern', ' — HexHive Solutions', 'Jan 2026 | Puducherry, India')
add_bullet('Programmed a Java Spring Boot backend connecting to AWS EC2 and S3 instances to process and store incoming IoT sensor data.')
add_bullet('Created REST APIs enabling HTTP communication between the hardware endpoints and the frontend dashboard, successfully deploying the software on cloud servers.')

# ========== CERTIFICATIONS ==========
add_section_header('Certifications')
cert_p = doc.add_paragraph()
cert_run = cert_p.add_run(
    'Java Programming  |  C Programming  |  Python Programming  |  '
    'Digital Circuits & Cloud Computing – NPTEL  |  Automation Developer Associate – ICT Academy'
)
cert_run.font.size = Pt(12)

# ========== LANGUAGES ==========
add_section_header('Languages')
lang_p = doc.add_paragraph()
lang_run = lang_p.add_run('English (Professional)  |  French (Beginner)  |  Tamil (Native)')
lang_run.font.size = Pt(12)

# ========== SAVE ==========
output_path = r'C:\Users\Asus\OneDrive\Desktop\resume\Siddharth_K_Resume_AntiGravity.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
