from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== ULTRA TIGHT PAGE MARGINS ==========
for section in doc.sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# ========== HELPER FUNCTIONS ==========
def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_entry_line(left_bold, left_normal, right_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.2), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    if left_bold:
        run_b = p.add_run(left_bold)
        run_b.font.size = Pt(12)
        run_b.font.bold = True
    if left_normal:
        run_n = p.add_run(left_normal)
        run_n.font.size = Pt(12)
        run_n.font.bold = False
    if right_text:
        run_t = p.add_run('\t' + right_text)
        run_t.font.size = Pt(12)
        run_t.font.bold = False

def add_sub_line(left_text, right_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.2), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(left_text)
    run.font.size = Pt(12)
    run.font.italic = True
    if right_text:
        run_r = p.add_run('\t' + right_text)
        run_r.font.size = Pt(12)
        run_r.font.italic = True

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run("• " + text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_normal_text(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(12)

# ========== NAME & CONTACT ==========
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(0)
name_run = name_p.add_run('SIDDHARTH K')
name_run.font.size = Pt(14)
name_run.font.bold = True
name_run.font.color.rgb = RGBColor(0, 0, 0)

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(2)
contact_run = contact_p.add_run('+91 9994475724  |  siddharthkadiresan21@gmail.com  |  LinkedIn: Siddharth K  |  GitHub: Siddharth-cype  |  Puducherry, India')
contact_run.font.size = Pt(10)

# ========== SUMMARY ==========
add_section_header('Professional Summary')
add_normal_text(
    'Pre-final B.Tech ECE student experienced in full-stack development, embedded systems, IoT architecture, and ML. '
    'Proficient in Java, Python, C/C++, and building distributed, cloud-connected systems using AWS IoT Core, '
    'Spring Boot, and ESP32. Seeking roles in embedded engineering, backend development, and cloud infrastructure.'
)

# ========== TECHNICAL SKILLS ==========
add_section_header('Technical Skills')
skills = [
    ('Languages: ', 'Java, Python, C, C++, Embedded C, HTML5, CSS3, JavaScript, SQL'),
    ('Embedded & IoT: ', 'ESP32, Sensor Interfacing, MQTT, mTLS, Real-Time Data Acquisition, PWM Control'),
    ('Cloud & Backend: ', 'AWS IoT Core, AWS EC2, AWS S3, Spring Boot, REST APIs, Microservices Architecture'),
    ('ML & AI: ', 'TensorFlow, Keras, YOLO, Vision Transformer (ViT), OpenCV, MediaPipe, Image Classification'),
    ('Tools & Databases: ', 'Git, GitHub, VS Code, Arduino IDE, Docker, Figma | MySQL, JPA/Hibernate')
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run_label = p.add_run(label)
    run_label.font.size = Pt(12)
    run_label.font.bold = True
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)

# ========== EXPERIENCE ==========
add_section_header('Experience')
add_entry_line('Software Engineering Intern', '  —  HexHive Solutions', 'Jan 2026')
add_sub_line('Puducherry, India', '')
add_bullet('Engineered a Spring Boot backend integrated with AWS EC2 & S3 for real-time IoT telemetry ingestion and storage.')
add_bullet('Designed scalable REST APIs for bidirectional hardware-software communication supporting low-latency MQTT data pipelines.')

# ========== PROJECTS ==========
add_section_header('Projects')
add_entry_line('Hexive — AWS IoT Environmental Monitoring System', '', 'Jan 2026')
add_sub_line('ESP32, AWS IoT Core, MQTT, mTLS, Spring Boot, Java, JPA', '')
add_bullet('Architected a distributed IoT ecosystem with ESP32 edge nodes publishing multi-sensor telemetry to AWS IoT Core via mTLS.')
add_bullet('Built a priority-cascading Java Decision Engine classifying environmental threats with real-time dashboard visualization.')

add_entry_line('Vision Transformer (ViT) for Image Classification', '', 'Dec 2025')
add_sub_line('Python, TensorFlow, Keras, Deep Learning', '')
add_bullet('Implemented a Vision Transformer with patch embedding and multi-head attention; trained on CIFAR-10 with competitive accuracy.')

add_entry_line('Computer Vision Touchless System Control', '', 'May 2025')
add_sub_line('Python, OpenCV, MediaPipe, Gesture Recognition', '')
add_bullet('Developed real-time hand gesture recognition system mapping gestures to cursor movement, clicks, scrolling, and screenshots.')

add_entry_line('Smart IoT Automated Irrigation System', '', 'May 2024')
add_sub_line('ESP32, C/C++, MicroPython, Cloud Dashboard', '')
add_bullet('Designed precision irrigation system with real-time moisture monitoring, cloud dashboard, and threshold-based actuation.')

# ========== EDUCATION ==========
add_section_header('Education')
add_entry_line('Manakula Vinayagar Institute of Technology', '', 'Sep 2023 – Apr 2027')
add_sub_line('B.Tech — Electronics and Communication Engineering  |  CGPA: 8.1', 'Puducherry')
add_entry_line('Vivekananda Higher Secondary School', '', 'Mar 2023')
add_sub_line('Higher Secondary (Class XII)  |  77.3%', 'Puducherry')

# ========== CERTIFICATIONS & LANGUAGES ==========
add_section_header('Certifications & Languages')
add_normal_text('Certifications: Java, Python, C, HTML/CSS/JS  |  Digital Circuits & Cloud (NPTEL)  |  Automation Assoc (ICT Academy)')
add_normal_text('Languages: English (Professional), French (Beginner), Tamil (Native)')

# ========== SAVE ==========
output_path = r'C:\Users\Asus\OneDrive\Desktop\resume\Siddharth_K_Resume_ATS_1Page.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
