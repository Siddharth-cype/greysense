from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ========== ULTRA TIGHT PAGE MARGINS FOR 1-PAGE FIT ==========
for section in doc.sections:
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = RGBColor(0, 0, 0)
pf = style.paragraph_format
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.line_spacing = 1.0

# ========== HELPER FUNCTIONS ==========
def add_section_header(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_entry_line(left_bold, left_normal, right_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    if left_bold:
        run_b = p.add_run(left_bold)
        run_b.font.size = Pt(12)
        run_b.font.bold = True
    if left_normal:
        run_n = p.add_run(left_normal)
        run_n.font.size = Pt(12)
        run_n.font.bold = False
    if right_text:
        run_t = p.add_run('\t' + right_text)
        run_t.font.size = Pt(12)
        run_t.font.bold = False

def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(f"• {text}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

# ========== NAME & CONTACT ==========
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(0)
name_run = name_p.add_run('SIDDHARTH K')
name_run.font.size = Pt(14)
name_run.font.bold = True

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(2)
contact_run = contact_p.add_run('+91 9994475724  |  siddharthkadiresan21@gmail.com  |  LinkedIn: Siddharth K  |  GitHub: Siddharth-cype')
contact_run.font.size = Pt(12)

# ========== SUMMARY ==========
add_section_header('Professional Summary')
summary_p = doc.add_paragraph()
summary_p.paragraph_format.space_after = Pt(0)
summary_run = summary_p.add_run(
    'B.Tech ECE engineering student specializing in Embedded Systems and scalable backend architectures. '
    'Proven ability to orchestrate Edge Computing hardware and low-latency cloud data pipelines using AWS IoT Core and Spring Boot. '
    'Experienced in translating raw sensor telemetry and Neural Networks outputs into real-time business logic. '
    'Actively seeking roles in Firmware Development, systems architecture, and distributed IoT applications.'
)
summary_run.font.size = Pt(12)

# ========== EDUCATION ==========
add_section_header('Education')
add_entry_line('Manakula Vinayagar Institute of Technology', '', 'Sep 2023 – Apr 2027')
add_entry_line('B.Tech — Electronics and Communication Engineering (CGPA: 8.1)', '', 'Puducherry, India')

add_entry_line('Vivekananda Higher Secondary School', '', 'Mar 2023')
add_entry_line('Higher Secondary (Class XII) — 77.3%', '', 'Puducherry, India')

# ========== TECHNICAL SKILLS ==========
add_section_header('Technical Skills')
skills = [
    ('Languages: ', 'Java, Python, C, C++, Embedded C, SQL, HTML5, CSS3, JavaScript'),
    ('Embedded Systems & IoT: ', 'ESP32, Firmware Development, Edge Computing, Sensor Interfacing, MQTT, mTLS, PWM'),
    ('Cloud & Backend: ', 'AWS IoT Core, AWS EC2, AWS S3, Spring Boot, REST APIs, Microservices, Real-time Streaming'),
    ('ML & Advanced Tech: ', 'Neural Networks, Vision Transformer (ViT), TensorFlow, Keras, OpenCV, MediaPipe, YOLO'),
    ('Tools & Databases: ', 'Git, GitHub, Docker, Maven, VS Code, Arduino IDE  |  MySQL, JPA/Hibernate, H2')
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run_label = p.add_run(label)
    run_label.font.size = Pt(12)
    run_label.font.bold = True
    run_value = p.add_run(value)
    run_value.font.size = Pt(12)

# ========== PROJECTS ==========
add_section_header('Projects')

add_entry_line('Hexive — Distributed IoT Environmental Monitoring System', '', 'Jan 2026')
add_entry_line('Tech Stack: ESP32, AWS IoT Core, Java Spring Boot, MQTT, mTLS, JPA', '', 'Puducherry, India')
add_bullet('Engineered a highly scalable edge-to-cloud telemetry pipeline by integrating ESP32 nodes with AWS IoT Core over mTLS, achieving sub-100ms low-latency data transmission for 4 parallel sensor streams.')
add_bullet('Developed a robust Spring Boot backend handling 100+ concurrent state aggregations, driving a real-time analytics dashboard used for proactive facility monitoring and threat detection.')

add_entry_line('HemoSight Pro — AI-Powered Conjunctival Hemoglobin Estimation', '', 'Feb 2026')
add_entry_line('Tech Stack: Python, OpenCV, MediaPipe, Haar Cascades, Flask', '', 'Puducherry, India')
add_bullet('Architected a continuous ROI tracking system executing real-time image quality validation using Laplacian variance, reducing false-positive AI predictions under extreme camera proximity constraints.')
add_bullet('Deployed a triple-layer computer vision pipeline combining MediaPipe Mesh and Haar Cascades, resulting in highly stable physiological data extraction suitable for clinical-grade hardware enclosures.')

add_entry_line('Vision Transformer (ViT) Image Classifier', '', 'Dec 2025')
add_entry_line('Tech Stack: Python, TensorFlow, Keras, Neural Networks', '', 'Puducherry, India')
add_bullet('Implemented a complex Vision Transformer architecture entirely from scratch (including patch embedding and multi-head attention), demonstrating advanced comprehension of modern Neural Networks and deep learning principles.')

# ========== INTERNSHIP ==========
add_section_header('Experience')
add_entry_line('Software Engineering Intern  —  HexHive Solutions', '', 'Jan 2026')
add_entry_line('Focus: Backend Architecture & Cloud Infrastructure', '', 'Puducherry, India')
add_bullet('Led Firmware Development protocols bridging ESP32 microcontrollers to cloud endpoints, executing highly secure data handoffs using industry-standard cryptography.')
add_bullet('Accelerated deployment velocity by containerizing Edge Computing frameworks and orchestrating EC2/S3 resources, resulting in a 30% aggregate decrease in API response latency.')

# ========== CERTIFICATIONS ==========
add_section_header('Certifications')
cert_p = doc.add_paragraph()
cert_p.paragraph_format.space_after = Pt(0)
cert_run = cert_p.add_run(
    'Digital Circuits & Cloud Computing (NPTEL)  |  Automation Developer Associate (ICT Academy)  |  '
    'Java, Python, C Programming  |  RPA Design & Development'
)
cert_run.font.size = Pt(12)

# ========== LANGUAGES ==========
add_section_header('Languages')
lang_p = doc.add_paragraph()
lang_p.paragraph_format.space_after = Pt(0)
lang_run = lang_p.add_run('English (Professional)  |  French (Beginner)  |  Tamil (Native)')
lang_run.font.size = Pt(12)

# ========== SAVE ==========
output_path = r'C:\Users\Asus\OneDrive\Desktop\resume\Siddharth_K_Resume_Strict_ATS.docx'
doc.save(output_path)
print(f'Resume saved to: {output_path}')
